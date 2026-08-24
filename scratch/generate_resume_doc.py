import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_resume_doc(filepath):
    doc = Document()
    
    # Adjust margins to fit well on 2 pages
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Title
    title = doc.add_heading('Intelligent Traffic Control System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Overview
    doc.add_heading('1. Project Overview', level=1)
    p = doc.add_paragraph(
        "The Intelligent Traffic Control System is an advanced, computer vision-based platform designed to optimize traffic light timings dynamically. Traditional traffic lights operate on fixed timers, often leading to unnecessary congestion, longer wait times, and increased carbon emissions. This project solves this problem by utilizing real-time video feeds from traffic cameras to estimate vehicle density and intelligently allocate green-light times for different lanes, thereby ensuring smoother traffic flow."
    )
    
    # Objectives
    doc.add_heading('2. Key Objectives', level=1)
    doc.add_paragraph("Dynamic Traffic Management: Replace static timers with a dynamic time allocation algorithm based on real-time vehicle counts.", style='List Bullet')
    doc.add_paragraph("Real-Time Vehicle Detection: Accurately identify and classify various types of vehicles (cars, motorcycles, buses, trucks) using state-of-the-art Deep Learning models.", style='List Bullet')
    doc.add_paragraph("Live Dashboard: Provide a user-friendly frontend dashboard for monitoring traffic states, densities, and current active lanes in real-time.", style='List Bullet')

    # Tech Stack
    doc.add_heading('3. Technologies and Tools Used', level=1)
    doc.add_paragraph("Machine Learning & Computer Vision: YOLOv8 (You Only Look Once version 8), OpenCV, PyTorch.", style='List Bullet')
    doc.add_paragraph("Backend Framework: Python, Flask, Server-Sent Events (SSE).", style='List Bullet')
    doc.add_paragraph("Frontend: HTML5, CSS3, JavaScript (Vanilla JS with async fetching).", style='List Bullet')
    doc.add_paragraph("Data Handling: NumPy, JSON.", style='List Bullet')

    # Architecture & Implementation
    doc.add_heading('4. System Architecture & Implementation', level=1)
    
    doc.add_heading('4.1 Computer Vision & Vehicle Detection', level=2)
    p = doc.add_paragraph(
        "The core of the detection system is built upon the YOLOv8 object detection model. YOLOv8 was chosen for its high accuracy and rapid inference speed, which is critical for real-time applications. "
        "The system processes incoming image frames (resized to 640x480 for consistent processing) and filters predictions to identify specific vehicle classes. Bound boxes are drawn around detected vehicles, and the system aggregates a total vehicle count for the given frame."
    )
    
    doc.add_heading('4.2 Dynamic Timing Algorithm', level=2)
    p = doc.add_paragraph(
        "Once the vehicle count is obtained, a heuristic-based machine learning logic layer translates the vehicle count into a predicted green-light time. The timer logic operates in increments:"
    )
    doc.add_paragraph("0 vehicles: 0 seconds (Skip to next lane).", style='List Bullet')
    doc.add_paragraph("1-10 vehicles: 15 seconds (LOW Density).", style='List Bullet')
    doc.add_paragraph("11-20 vehicles: 30 seconds (LOW-MEDIUM Density).", style='List Bullet')
    doc.add_paragraph("For every 10 additional vehicles, the green time is extended by 15 seconds, maxing out at 120 seconds (SEVERE Density).", style='List Bullet')

    doc.add_heading('4.3 Backend & Synchronization', level=2)
    p = doc.add_paragraph(
        "The backend is developed in Flask. It manages the states of multiple camera nodes (e.g., North, South, East, West). A background thread continuously runs the traffic controller, transitioning active lanes based on the allocated green time and current queue state. "
        "The backend provides RESTful endpoints to upload new frames, fetch video feeds, and an SSE (Server-Sent Events) endpoint to push continuous real-time state updates (counts, densities, active camera) to the frontend."
    )
    
    # Page Break for formatting
    doc.add_page_break()

    # Results & Impact
    doc.add_heading('5. Results & Impact', level=1)
    p = doc.add_paragraph(
        "The deployment of this intelligent system drastically reduces idle times at intersections. By dynamically allocating zero seconds to empty lanes, the system avoids the common pitfall of halting traffic for a non-existent cross-flow. The dashboard provides immediate visual feedback, allowing operators to monitor the health and status of the intersection efficiently. "
        "The integration of YOLOv8 ensures robustness against varying lighting and weather conditions, while the SSE-based dashboard guarantees low-latency monitoring."
    )

    # My Role & Learnings
    doc.add_heading('6. My Role & Key Learnings', level=1)
    p = doc.add_paragraph(
        "As the lead developer on this project, I was responsible for the end-to-end implementation. Key learnings include:"
    )
    doc.add_paragraph("Deep Learning Integration: Successfully integrating a PyTorch-based YOLOv8 model into a lightweight web server.", style='List Bullet')
    doc.add_paragraph("Concurrency in Python: Managing background threads for the traffic controller alongside a Flask web server serving synchronous and asynchronous requests.", style='List Bullet')
    doc.add_paragraph("Real-Time Web Communication: Replacing traditional polling with Server-Sent Events (SSE) to push state updates efficiently to the client side.", style='List Bullet')
    doc.add_paragraph("System Optimization: Balancing the trade-off between image resolution, inference speed, and detection accuracy to achieve real-time performance.", style='List Bullet')

    # Future Enhancements
    doc.add_heading('7. Future Enhancements', level=1)
    doc.add_paragraph("Emergency Vehicle Override: Implementing audio/visual detection for sirens and emergency lights to instantly grant right-of-way.", style='List Bullet')
    doc.add_paragraph("Reinforcement Learning: Replacing the heuristic timer logic with a Reinforcement Learning agent (e.g., Deep Q-Network) that learns optimal timings from continuous simulation and feedback.", style='List Bullet')
    doc.add_paragraph("Multi-Intersection Synchronization: Scaling the system to communicate with adjacent intersections, creating a 'green wave' effect across city corridors.", style='List Bullet')

    doc.save(filepath)
    print(f"Document saved to {filepath}")

if __name__ == "__main__":
    filepath = r"d:\projects\New_ML_Project\Intelligent_Traffic_Control_System_Resume_Details.docx"
    create_resume_doc(filepath)
